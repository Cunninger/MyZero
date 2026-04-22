from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks, UploadFile, File, Form
from sqlalchemy.orm import Session
from fastapi.responses import StreamingResponse, JSONResponse
from typing import List, Dict, Optional
import asyncio
import threading
import time
import json
import io

from app.database import get_db, SessionLocal
from app.models import OptimizationRecord, OptimizationSegment, ChangeLog, AppConfig
from app.schemas import (
    TextOptimizeRequest, TextOptimizeResponse,
    SegmentResponse, ChangeLogResponse, RecordDetailResponse,
    ProgressUpdate
)
from app.services.ai_service import ai_service, count_text_length, SEGMENT_SKIP_THRESHOLD
from app.services.mineru_service import mineru_service, needs_mineru

router = APIRouter(prefix="/optimize", tags=["optimization"])

# Global SSE message storage (record_id -> list of messages)
sse_messages: Dict[int, List[dict]] = {}
sse_lock = threading.Lock()
sse_active: Dict[int, bool] = {}

# Track cancelled records
cancelled_records: set = set()
cancelled_lock = threading.Lock()

COMPRESSION_THRESHOLD = 5000


def process_mineru_upload(record_id: int, file_bytes: bytes, filename: str, mode: str):
    """Background task: parse file via MinerU, then optimize."""
    db = SessionLocal()
    try:
        record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
        if not record:
            return

        sse_active[record_id] = True
        record.status = "parsing"
        db.commit()

        def on_progress(msg):
            add_sse_message(record_id, {"type": "parsing", "message": msg})

        # Parse via MinerU
        try:
            text = mineru_service.parse_file(file_bytes, filename, on_progress=on_progress)
        except Exception as e:
            record.status = "failed"
            record.error_message = f"文档解析失败: {str(e)[:200]}"
            db.commit()
            add_sse_message(record_id, {"type": "failed", "error": f"文档解析失败: {str(e)[:200]}"})
            cleanup_sse(record_id)
            db.close()
            return

        # Update record with parsed text
        record.original_text = text
        db.commit()
        add_sse_message(record_id, {"type": "parsing_done", "text_length": len(text)})

        # Close this DB session before starting optimization (process_optimization opens its own)
        db.close()

        # Run the optimization pipeline
        process_optimization(record_id, text, mode)

    except Exception as e:
        try:
            record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
            if record:
                record.status = "failed"
                record.error_message = str(e)[:500]
                db.commit()
        except Exception:
            pass
        add_sse_message(record_id, {"type": "failed", "error": str(e)[:200]})
        cleanup_sse(record_id)
    finally:
        try:
            db.close()
        except Exception:
            pass


def add_sse_message(record_id: int, data: dict):
    with sse_lock:
        if record_id not in sse_messages:
            sse_messages[record_id] = []
        sse_messages[record_id].append(data)


def cleanup_sse(record_id: int):
    with sse_lock:
        if record_id in sse_messages:
            del sse_messages[record_id]
        sse_active[record_id] = False


def _get_api_interval() -> int:
    """Get API request interval from DB config."""
    try:
        db = SessionLocal()
        config = db.query(AppConfig).filter(AppConfig.id == 1).first()
        db.close()
        if config and config.api_request_interval:
            return config.api_request_interval
    except Exception:
        pass
    return 0


def process_optimization(record_id: int, text: str, mode: str):
    """Background task: two-stage processing with history context, compression, streaming."""
    db = SessionLocal()
    try:
        record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
        if not record:
            return

        sse_active[record_id] = True
        record.status = "processing"
        db.commit()

        api_interval = _get_api_interval()

        # Split text
        segments_data = ai_service.split_text_into_segments(text)
        total_segments = len(segments_data)
        record.total_segments = total_segments
        db.commit()

        # Create segment rows
        for idx, seg_info in enumerate(segments_data):
            seg = OptimizationSegment(
                record_id=record_id,
                segment_index=idx,
                stage="pending",
                original_text=seg_info["text"],
                status="pending",
            )
            db.add(seg)
        db.commit()

        # Determine stages
        if mode == "combined":
            stages = ["polish", "enhance"]
        elif mode == "polish":
            stages = ["polish"]
        elif mode == "humanize":
            stages = ["enhance"]
        else:
            stages = [mode]

        total_stages = len(stages)

        # Process each stage
        for stage_idx, stage in enumerate(stages):
            history: List[Dict] = []
            total_chars = 0

            db_segments = db.query(OptimizationSegment).filter(
                OptimizationSegment.record_id == record_id
            ).order_by(OptimizationSegment.segment_index).all()

            for seg_idx, segment in enumerate(db_segments):
                # Check cancellation
                with cancelled_lock:
                    if record_id in cancelled_records:
                        record.status = "stopped"
                        record.error_message = "用户取消"
                        db.commit()
                        add_sse_message(record_id, {
                            "type": "stopped",
                            "segment_index": seg_idx,
                            "total_segments": total_segments,
                        })
                        cleanup_sse(record_id)
                        cancelled_records.discard(record_id)
                        return

                # Calculate overall progress
                if total_stages > 1:
                    progress = (stage_idx + seg_idx / total_segments) / total_stages
                else:
                    progress = seg_idx / total_segments if total_segments > 0 else 0

                # Update segment stage
                segment.stage = stage
                db.commit()

                # Skip short segments (titles)
                if count_text_length(segment.original_text) < SEGMENT_SKIP_THRESHOLD:
                    if stage == "polish":
                        segment.polished_text = segment.original_text
                    segment.optimized_text = segment.original_text
                    segment.status = "completed"
                    from datetime import datetime
                    segment.completed_at = datetime.now()
                    db.commit()

                    add_sse_message(record_id, {
                        "type": "segment_complete",
                        "segment_index": seg_idx,
                        "total_segments": total_segments,
                        "stage": stage,
                        "stage_index": stage_idx,
                        "total_stages": total_stages,
                        "progress": progress,
                    })
                    continue

                # Skip already-processed segments in later stages (for retry scenarios)
                if stage_idx > 0 and stage == "enhance" and segment.polished_text:
                    pass  # Will use polished_text as input below
                elif stage_idx > 0 and stage == "enhance" and not segment.polished_text:
                    # Polish stage didn't complete, use original
                    pass

                # Determine input text
                if stage == "enhance" and segment.polished_text:
                    input_text = segment.polished_text
                else:
                    input_text = segment.original_text

                # Streaming chunk callback
                chunk_buffer = {"text": ""}

                def on_chunk(chunk, _buf=chunk_buffer, _rid=record_id, _sidx=seg_idx, _stg=stage):
                    _buf["text"] += chunk
                    add_sse_message(_rid, {
                        "type": "content_chunk",
                        "segment_index": _sidx,
                        "stage": _stg,
                        "chunk": chunk,
                    })

                segment.status = "processing"
                db.commit()

                try:
                    result = ai_service.optimize_segment_with_history_sync(
                        input_text, stage, history, on_chunk=on_chunk
                    )

                    # Store result
                    if stage == "polish":
                        segment.polished_text = result
                    segment.optimized_text = result
                    segment.status = "completed"
                    from datetime import datetime
                    segment.completed_at = datetime.now()
                    db.commit()

                    # Record change log
                    changes_detail = ai_service.compute_changes_detail(input_text, result)
                    existing_change = db.query(ChangeLog).filter(
                        ChangeLog.record_id == record_id,
                        ChangeLog.segment_index == seg_idx,
                        ChangeLog.stage == stage,
                    ).first()
                    if existing_change:
                        existing_change.before_text = input_text
                        existing_change.after_text = result
                        existing_change.changes_detail = json.dumps(changes_detail, ensure_ascii=False)
                    else:
                        db.add(ChangeLog(
                            record_id=record_id,
                            segment_index=seg_idx,
                            stage=stage,
                            before_text=input_text,
                            after_text=result,
                            changes_detail=json.dumps(changes_detail, ensure_ascii=False),
                        ))
                    db.commit()

                    # Update history
                    history.append({"role": "assistant", "content": result})
                    total_chars += len(result)

                    # Compress history if needed
                    if total_chars > COMPRESSION_THRESHOLD:
                        history = ai_service.compress_history_sync(history, stage)
                        total_chars = sum(len(m.get("content", "")) for m in history)
                        add_sse_message(record_id, {"type": "history_compressed", "stage": stage})

                    add_sse_message(record_id, {
                        "type": "segment_complete",
                        "segment_index": seg_idx,
                        "total_segments": total_segments,
                        "stage": stage,
                        "stage_index": stage_idx,
                        "total_stages": total_stages,
                        "progress": progress,
                    })

                    # API interval
                    if api_interval > 0 and seg_idx < len(db_segments) - 1:
                        time.sleep(api_interval)

                except Exception as seg_error:
                    segment.status = "failed"
                    if not segment.optimized_text:
                        segment.optimized_text = input_text
                    db.commit()

                    add_sse_message(record_id, {
                        "type": "segment_complete",
                        "segment_index": seg_idx,
                        "total_segments": total_segments,
                        "stage": stage,
                        "stage_index": stage_idx,
                        "total_stages": total_stages,
                        "status": "failed",
                    })

        # Final assembly
        final_segments = db.query(OptimizationSegment).filter(
            OptimizationSegment.record_id == record_id
        ).order_by(OptimizationSegment.segment_index).all()

        optimized_text = '\n\n'.join(seg.optimized_text or seg.original_text for seg in final_segments)
        record.optimized_text = optimized_text
        record.status = "completed"
        from datetime import datetime
        record.completed_at = datetime.now()
        db.commit()

        add_sse_message(record_id, {
            "type": "completed",
            "total_segments": total_segments,
            "optimized_text": optimized_text,
            "record_id": record_id,
            "mode": record.mode,
            "original_text": record.original_text,
        })
        cleanup_sse(record_id)

    except Exception as e:
        try:
            record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
            if record:
                record.status = "failed"
                record.error_message = str(e)[:500]
                db.commit()
        except Exception:
            pass
        add_sse_message(record_id, {"type": "failed", "error": str(e)[:200]})
        cleanup_sse(record_id)
    finally:
        db.close()


@router.post("", response_model=TextOptimizeResponse)
async def create_optimization(
    request: TextOptimizeRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
):
    """Submit a text optimization request."""
    record = OptimizationRecord(
        original_text=request.text,
        mode=request.mode,
        status="pending",
    )
    db.add(record)
    db.commit()
    db.refresh(record)
    background_tasks.add_task(process_optimization, record.id, request.text, request.mode)
    return record


@router.post("/upload", response_model=TextOptimizeResponse)
async def create_optimization_from_file(
    file: UploadFile = File(...),
    mode: str = Form("combined"),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db),
):
    """Submit optimization from uploaded file (.txt, .docx)."""
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Route MinerU-supported formats (PDF, images, PPT)
    if needs_mineru(filename):
        content = await file.read()
        record = OptimizationRecord(
            original_text=f"[文件上传中: {filename}]",
            mode=mode,
            status="parsing",
        )
        db.add(record)
        db.commit()
        db.refresh(record)
        background_tasks.add_task(process_mineru_upload, record.id, content, filename, mode)
        return record

    # Existing text/DOCX handling
    if ext not in ("txt", "docx", "md", "markdown"):
        raise HTTPException(status_code=400, detail="不支持的文件格式，请上传 .txt .docx .pdf .ppt .pptx 或图片文件")

    content = await file.read()

    if ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            text = "\n".join(paragraphs)
        except ImportError:
            raise HTTPException(status_code=500, detail="服务端未安装 python-docx，无法解析 .docx 文件")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"解析 .docx 文件失败: {str(e)}")
    else:
        # txt, md
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                text = content.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            text = content.decode("utf-8", errors="replace")

    if not text.strip():
        raise HTTPException(status_code=400, detail="文件内容为空")

    record = OptimizationRecord(
        original_text=text,
        mode=mode,
        status="pending",
    )
    db.add(record)
    db.commit()
    db.refresh(record)
    background_tasks.add_task(process_optimization, record.id, text, mode)
    return record


@router.get("/{record_id}", response_model=TextOptimizeResponse)
async def get_optimization(record_id: int, db: Session = Depends(get_db)):
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")
    return record


@router.post("/{record_id}/stop")
async def stop_optimization(record_id: int, db: Session = Depends(get_db)):
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")
    if record.status not in ("pending", "processing"):
        raise HTTPException(status_code=400, detail="Cannot stop a completed or failed record")
    with cancelled_lock:
        cancelled_records.add(record_id)
    return {"message": "Stop signal sent", "record_id": record_id}


@router.post("/{record_id}/retry")
async def retry_failed_segments(record_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")

    failed_segments = db.query(OptimizationSegment).filter(
        OptimizationSegment.record_id == record_id,
        OptimizationSegment.status == "failed",
    ).all()

    if not failed_segments:
        raise HTTPException(status_code=400, detail="No failed segments to retry")

    record.status = "processing"
    db.commit()
    sse_active[record_id] = True

    def retry_segments():
        db_retry = SessionLocal()
        try:
            # Re-determine stages
            mode = record.mode
            if mode == "combined":
                stages = ["polish", "enhance"]
            elif mode == "polish":
                stages = ["polish"]
            else:
                stages = ["enhance"]

            api_interval = _get_api_interval()

            for seg in failed_segments:
                seg_db = db_retry.query(OptimizationSegment).filter(OptimizationSegment.id == seg.id).first()
                if not seg_db:
                    continue

                with cancelled_lock:
                    if record_id in cancelled_records:
                        rec = db_retry.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
                        if rec:
                            rec.status = "stopped"
                            db_retry.commit()
                        cancelled_records.discard(record_id)
                        return

                seg_db.status = "processing"
                db_retry.commit()

                # Determine which stage to retry
                stage = stages[-1] if seg_db.stage in stages else stages[0]
                input_text = seg_db.polished_text if (stage == "enhance" and seg_db.polished_text) else seg_db.original_text

                try:
                    optimized_text = ai_service.optimize_segment_with_history_sync(input_text, stage)
                    from datetime import datetime
                    if stage == "polish":
                        seg_db.polished_text = optimized_text
                    seg_db.optimized_text = optimized_text
                    seg_db.status = "completed"
                    seg_db.completed_at = datetime.now()
                    db_retry.commit()

                    changes_detail = ai_service.compute_changes_detail(input_text, optimized_text)
                    existing = db_retry.query(ChangeLog).filter(
                        ChangeLog.record_id == record_id,
                        ChangeLog.segment_index == seg_db.segment_index,
                        ChangeLog.stage == stage,
                    ).first()
                    if existing:
                        existing.after_text = optimized_text
                        existing.changes_detail = json.dumps(changes_detail, ensure_ascii=False)
                    else:
                        db_retry.add(ChangeLog(
                            record_id=record_id,
                            segment_index=seg_db.segment_index,
                            stage=stage,
                            before_text=input_text,
                            after_text=optimized_text,
                            changes_detail=json.dumps(changes_detail, ensure_ascii=False),
                        ))
                    db_retry.commit()

                    add_sse_message(record_id, {
                        "type": "segment_complete",
                        "segment_index": seg_db.segment_index,
                        "total_segments": record.total_segments,
                        "stage": stage,
                    })

                    if api_interval > 0:
                        time.sleep(api_interval)

                except Exception:
                    seg_db.status = "failed"
                    db_retry.commit()

            # Rebuild full text
            all_segs = db_retry.query(OptimizationSegment).filter(
                OptimizationSegment.record_id == record_id
            ).order_by(OptimizationSegment.segment_index).all()

            from datetime import datetime
            rec = db_retry.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
            rec.optimized_text = '\n\n'.join(s.optimized_text or s.original_text for s in all_segs)
            rec.status = "completed"
            rec.completed_at = datetime.now()
            db_retry.commit()

            add_sse_message(record_id, {
                "type": "completed",
                "total_segments": rec.total_segments,
                "optimized_text": rec.optimized_text,
                "record_id": record_id,
                "mode": rec.mode,
                "original_text": rec.original_text,
            })
            cleanup_sse(record_id)

        except Exception as e:
            rec = db_retry.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
            if rec:
                rec.status = "failed"
                rec.error_message = str(e)[:500]
                db_retry.commit()
            add_sse_message(record_id, {"type": "failed", "error": str(e)[:200]})
            cleanup_sse(record_id)
        finally:
            db_retry.close()

    background_tasks.add_task(retry_segments)
    return {"message": "Retrying failed segments", "failed_count": len(failed_segments)}


@router.get("/{record_id}/detail", response_model=RecordDetailResponse)
async def get_optimization_detail(record_id: int, db: Session = Depends(get_db)):
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")

    return {
        "id": record.id,
        "original_text": record.original_text,
        "optimized_text": record.optimized_text,
        "mode": record.mode,
        "status": record.status,
        "total_segments": record.total_segments,
        "error_message": record.error_message,
        "created_at": record.created_at,
        "completed_at": record.completed_at,
        "segments": [
            {
                "id": seg.id,
                "segment_index": seg.segment_index,
                "stage": seg.stage,
                "original_text": seg.original_text,
                "polished_text": seg.polished_text,
                "optimized_text": seg.optimized_text,
                "status": seg.status,
                "created_at": seg.created_at,
                "completed_at": seg.completed_at,
            }
            for seg in record.segments
        ],
        "changes": [
            {
                "id": change.id,
                "segment_index": change.segment_index,
                "stage": change.stage,
                "before_text": change.before_text,
                "after_text": change.after_text,
                "changes_detail": json.loads(change.changes_detail) if change.changes_detail else None,
                "created_at": change.created_at,
            }
            for change in record.changes
        ],
    }


@router.get("/{record_id}/segments", response_model=List[SegmentResponse])
async def get_optimization_segments(record_id: int, db: Session = Depends(get_db)):
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")
    return db.query(OptimizationSegment).filter(
        OptimizationSegment.record_id == record_id
    ).order_by(OptimizationSegment.segment_index).all()


@router.get("/{record_id}/changes", response_model=List[ChangeLogResponse])
async def get_optimization_changes(record_id: int, db: Session = Depends(get_db)):
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")

    changes = db.query(ChangeLog).filter(
        ChangeLog.record_id == record_id
    ).order_by(ChangeLog.segment_index).all()

    return [
        {
            "id": c.id,
            "segment_index": c.segment_index,
            "stage": c.stage,
            "before_text": c.before_text,
            "after_text": c.after_text,
            "changes_detail": json.loads(c.changes_detail) if c.changes_detail else None,
            "created_at": c.created_at,
        }
        for c in changes
    ]


@router.get("/{record_id}/progress", response_model=ProgressUpdate)
async def get_optimization_progress(record_id: int, db: Session = Depends(get_db)):
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")

    # Determine stages from mode
    mode = record.mode
    if mode == "combined":
        stages = ["polish", "enhance"]
    elif mode == "polish":
        stages = ["polish"]
    elif mode == "humanize":
        stages = ["enhance"]
    else:
        stages = [mode]
    total_stages = len(stages)

    total = record.total_segments or 1

    completed_count = db.query(OptimizationSegment).filter(
        OptimizationSegment.record_id == record_id,
        OptimizationSegment.status == "completed",
    ).count()

    # Old-style global progress (for backwards compatibility)
    progress = (completed_count / total) * 100 if total > 0 else 0

    # Calculate stage-aware overall progress
    if record.status == "completed":
        overall_progress = 100.0
        stage_index = total_stages - 1
    else:
        processing_segment = db.query(OptimizationSegment).filter(
            OptimizationSegment.record_id == record_id,
            OptimizationSegment.status == "processing",
        ).order_by(OptimizationSegment.segment_index).first()

        if processing_segment and processing_segment.stage in stages:
            stage_index = stages.index(processing_segment.stage)
            seg_idx = processing_segment.segment_index
            overall_progress = (stage_index + seg_idx / total) / total_stages * 100
        else:
            last_completed = db.query(OptimizationSegment).filter(
                OptimizationSegment.record_id == record_id,
                OptimizationSegment.status == "completed",
            ).order_by(OptimizationSegment.segment_index.desc()).first()

            if last_completed and last_completed.stage in stages:
                stage_index = stages.index(last_completed.stage)
                seg_idx = last_completed.segment_index + 1
                if seg_idx >= total:
                    seg_idx = 0
                    stage_index = min(stage_index + 1, total_stages - 1)
                overall_progress = (stage_index + seg_idx / total) / total_stages * 100
            else:
                stage_index = 0
                overall_progress = 0.0

    return ProgressUpdate(
        record_id=record.id,
        status=record.status,
        progress=round(progress, 1),
        overall_progress=round(overall_progress, 1),
        current_position=completed_count,
        total_segments=total,
        stage_index=stage_index,
        total_stages=total_stages,
        error_message=record.error_message,
    )


@router.get("/{record_id}/stream")
async def stream_optimization_progress(record_id: int, db: Session = Depends(get_db)):
    """SSE stream for real-time progress updates."""
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")

    async def event_generator():
        last_index = 0
        max_wait_cycles = 1200  # 10 minutes (0.5s * 1200)
        wait_cycles = 0

        while wait_cycles < max_wait_cycles:
            messages = []
            with sse_lock:
                messages = sse_messages.get(record_id, [])

            while last_index < len(messages):
                msg = messages[last_index]
                yield f"data: {json.dumps(msg)}\n\n"
                last_index += 1
                wait_cycles = 0

                if msg.get("type") in ("completed", "failed", "stopped"):
                    return

            await asyncio.sleep(0.5)
            wait_cycles += 1

        yield f"data: {json.dumps({'type': 'timeout'})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/{record_id}/export")
async def export_optimization(record_id: int, db: Session = Depends(get_db)):
    """Export optimized text."""
    record = db.query(OptimizationRecord).filter(OptimizationRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")
    if record.status != "completed":
        raise HTTPException(status_code=400, detail="Record not yet completed")

    segments = db.query(OptimizationSegment).filter(
        OptimizationSegment.record_id == record_id
    ).order_by(OptimizationSegment.segment_index).all()

    final_text = '\n\n'.join(seg.optimized_text or seg.original_text for seg in segments)

    return JSONResponse({
        "content": final_text,
        "filename": f"myzero_optimized_{record_id}.txt",
        "format": "txt",
    })
